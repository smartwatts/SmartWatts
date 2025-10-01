#!/usr/bin/env python3
"""
Convert Markdown Installation Guide to Word Document
Simple script to create a downloadable .docx file
"""

import os
import sys
from pathlib import Path

def create_docx_from_markdown():
    """Convert the markdown guide to a Word document."""
    
    # Check if python-docx is available
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Installing required package...")
        os.system("pip install python-docx")
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    
    # Read the markdown file
    markdown_file = Path(__file__).parent / "SmartWatts_Edge_Gateway_Installation_Guide.md"
    
    if not markdown_file.exists():
        print(f"Error: {markdown_file} not found!")
        return False
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    title_style.font.size = Inches(0.3)
    
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.size = Inches(0.25)
    
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Inches(0.22)
    
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.size = Inches(0.2)
    
    # Parse markdown content
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Add empty paragraph
            doc.add_paragraph()
            continue
        
        # Title
        if line.startswith('# '):
            title = line[2:].strip()
            doc.add_heading(title, level=0)
        
        # Heading 1
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
        
        # Heading 2
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
        
        # Heading 3
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)
        
        # Code blocks
        elif line.startswith('```'):
            # Skip code block markers
            continue
        
        # Bullet points
        elif line.startswith('- '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = line[3:].strip()
            doc.add_paragraph(text, style='List Number')
        
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        
        # Check marks
        elif line.startswith('- ✅'):
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(f"✅ {text}")
            run.bold = True
        
        # Regular paragraphs
        else:
            if line:
                # Handle inline formatting
                paragraph = doc.add_paragraph()
                
                # Split by ** for bold text
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        # Regular text
                        if part:
                            paragraph.add_run(part)
                    else:
                        # Bold text
                        run = paragraph.add_run(part)
                        run.bold = True
    
    # Save document
    output_file = Path(__file__).parent / "SmartWatts_Edge_Gateway_Installation_Guide.docx"
    doc.save(str(output_file))
    
    print(f"✅ Word document created: {output_file}")
    print(f"📄 File size: {output_file.stat().st_size / 1024:.1f} KB")
    
    return True

if __name__ == "__main__":
    print("🔄 Converting Markdown to Word document...")
    
    if create_docx_from_markdown():
        print("✅ Conversion completed successfully!")
        print("\n📁 You can now download the .docx file from:")
        print("   edge-gateway/docs/SmartWatts_Edge_Gateway_Installation_Guide.docx")
    else:
        print("❌ Conversion failed!")
        sys.exit(1)
